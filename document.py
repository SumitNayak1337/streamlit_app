from docx import Document
from docx.shared import Inches
import io
import re

def create_docx_from_text(ai_response_text: str, images_dict: dict = None) -> io.BytesIO:
    """
    Parses Gemini Markdown output and constructs an in-memory DOCX file.
    Using io.BytesIO so Streamlit can serve it directly as a download
    without saving it to the server's disk.
    If Gemini outputs [IMAGE: filename.png], it inserts the image into the doc.
    """
    doc = Document()
    
    # Basic Markdown parser for the DOCX format
    for line in ai_response_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('[IMAGE:'):
            match = re.search(r'\[IMAGE:\s*(.+?)\]', line)
            if match and images_dict:
                filename = match.group(1).strip()
                if filename in images_dict:
                    images_dict[filename].seek(0)
                    try:
                        doc.add_picture(images_dict[filename], width=Inches(6.0))
                    except Exception as e:
                        doc.add_paragraph(f"[Error loading image {filename}: {e}]")
                else:
                    doc.add_paragraph(f"[Image not found in attachments: {filename}]")
            else:
                doc.add_paragraph(line)
        else:
            # Clean up basic bolding for standard paragraphs
            clean_text = line.replace('**', '')
            doc.add_paragraph(clean_text)
            
    # Save to memory stream instead of disk
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream